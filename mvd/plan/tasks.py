# from background_task import background
from django.shortcuts import render, redirect, get_object_or_404
from plan.models import Mesyac, Profile, Kafedra, Plan, Predmet, NIR, VR, DR, UMR, INR, Nagruzka, DocInfo
from plan.forms import ShapkaForm, Table1Form, Table2Form, Table3Form, Table4Form, Table6Form, Table5Form, MainTableForm, Table1UploadForm, NagruzkaForm
from rating.models import URR, ORMR, PCR, MRR, Rating
from django.forms.formsets import formset_factory
from django.forms.models import modelformset_factory
from plan.Parser_and_overview import createDoc2, createDoc, takeTable, takeXls, writeInfoDoc, xlsPrepod
from django.core.files.base import ContentFile
from django.core.files import File
from django.contrib.auth.models import User
from django.conf import settings
from django.http import HttpResponse
#from plan.itogKafDoc import createDocument
#from plan.RatingAllKafDoc import writeAllKafDoc
import random
from docx import Document
import os


# exec(open('plan/tasks.py', encoding='utf-8').read())
"""назначения порядка месяцам"""


def mess_order():
    mes_list = Mesyac.objects.all()
    for mes in mes_list:
        if mes.name == "АВГУСТ":
            mes.order = 1
        if mes.name == "СЕНТЯБРЬ":
            mes.order = 2
        if mes.name == "ОКТЯБРЬ":
            mes.order = 3
        if mes.name == "НОЯБРЬ":
            mes.order = 4
        if mes.name == "ДЕКАБРЬ":
            mes.order = 5
        if mes.name == "Итого за 1 полугодие:":
            mes.order = 6
        if mes.name == "ЯНВАРЬ":
            mes.order = 7
        if mes.name == "ФЕВРАЛЬ":
            mes.order = 8
        if mes.name == "МАРТ":
            mes.order = 9
        if mes.name == "АПРЕЛЬ":
            mes.order = 10
        if mes.name == "МАЙ":
            mes.order = 11
        if mes.name == "ИЮНЬ":
            mes.order = 12
        if mes.name == "ИЮЛЬ":
            mes.order = 13
        if mes.name == "Итого за 2 полугодие:":
            mes.order = 14
        if mes.name == "Итого за учебный год:":
            mes.order = 15
        mes.save()


"""Пересчет сумм и эфф ставок всех сотрудников"""


def set_effect_summ(year):
    allrating = Rating.objects.all()
    count_success = 0
    count_deny = 0
    for r in allrating:
         # считаем общие позиции
        try:
            urrsumm = (URR.objects.get(profile=r.profile, year=year)).getsumm()
        except:
            urrsumm = 0
        try:
            ormrsumm = (ORMR.objects.get(
                profile=r.profile, year=year)).getsumm()
        except:
            ormrsumm = 0
        try:
            pcrsumm = (PCR.objects.get(profile=r.profile, year=year)).getsumm()
        except:
            pcrsumm = 0
        mrrsumm = 0
        mrrs = MRR.objects.filter(profile=r.profile, year=year)
        for m in mrrs:
            mrrsumm += m.bal
        r.summ = (urrsumm + ormrsumm + pcrsumm + mrrsumm)
        r.urr = urrsumm
        r.ormr = ormrsumm
        r.pcr = pcrsumm
        r.mrr = mrrsumm
        if r.kolvomes > 11:
            r.kolvomes = 11
        try:
            effect_stavka = (r.profile.info.stavka*r.kolvomes/11)
            r.effect_stavka = effect_stavka
            r.summ_effect = (urrsumm + ormrsumm +
                             pcrsumm + mrrsumm)/effect_stavka
            r.save()
            count_success += 1
        except:
            r.summ_effect = 0
            r.effect_stavka = -1
            r.save()
            count_deny += 1
    print(str(count_success)+" успешно заменены суммы ")
    print(str(count_deny)+" ошибка заменены суммы ")


"""Пересчет рейтинга для всех сотрудников"""


def refresh_places(year):
    set_effect_summ(year)
    # рейтинг с учетом ставки более 0,8
    allrating = Rating.objects.filter(year=year).order_by(
        "-summ_effect").exclude(effect_stavka__lte=0.8)
    # рейтинг полный-виртуальный
    allrating_effect = Rating.objects.filter(
        year=year).order_by("-summ_effect")

    kafedras = Kafedra.objects.all()
    dolzhnost_list = ['профессор', 'доцент', 'ст.преподаватель',
        'нач.кафедры', 'зам.нач.кафедры', 'преподаватель', ]

    # рейтинг по унику для всех
    count_places = 1
    for i in range(len(allrating)-1):
        print(allrating[i].profile.fullname)
        if allrating[i].summ_effect != allrating[i+1].summ_effect:
            allrating[i].unikplace = count_places
            count_places += 1
        else:
            allrating[i].unikplace = count_places
        allrating[i].save()
    if allrating[len(allrating)-1].summ_effect != allrating[len(allrating)-2].summ_effect:
        allrating[len(allrating)-1].unikplace = count_places
        allrating[len(allrating)-1].save()
    else:
        allrating[len(allrating)-1].unikplace = count_places-1
        allrating[len(allrating)-1].save()

    count_places = 1
    for i in range(len(allrating_effect)-1):
        print(allrating_effect[i].profile.fullname)
        if allrating_effect[i].summ_effect != allrating_effect[i+1].summ_effect:
            allrating_effect[i].unikplace_effect = count_places
            count_places += 1
        else:
            allrating_effect[i].unikplace_effect = count_places
        allrating_effect[i].save()
    if allrating_effect[len(allrating_effect)-1].summ_effect != allrating_effect[len(allrating_effect)-2].summ_effect:
        allrating_effect[len(allrating_effect) -
                             1].unikplace_effect = count_places
        allrating_effect[len(allrating_effect)-1].save()
    else:
        allrating_effect[len(allrating_effect) -
                             1].unikplace_effect = count_places-1
        allrating_effect[len(allrating_effect)-1].save()

    # рейтинг по кафедре для всех сотрудников
    for kafedra in kafedras:
        count_places = 1
        kafedra_rating_list = Rating.objects.filter(year=year, profile__kafedra=kafedra).order_by(
            "-summ_effect").exclude(effect_stavka__lte=0.8)
        for i in range(len(kafedra_rating_list)-1):
            if kafedra_rating_list[i].summ_effect != kafedra_rating_list[i+1].summ_effect:
                kafedra_rating_list[i].kafedraplace = count_places
                count_places += 1
            else:
                kafedra_rating_list[i].kafedraplace = count_places
            kafedra_rating_list[i].save()
        try:
            if kafedra_rating_list[len(kafedra_rating_list)-1].summ_effect != kafedra_rating_list[len(kafedra_rating_list)-2].summ_effect:
                kafedra_rating_list[len(
                    kafedra_rating_list)-1].kafedraplace = count_places
            else:
                kafedra_rating_list[len(
                    kafedra_rating_list)-1].kafedraplace = count_places - 1
            kafedra_rating_list[len(kafedra_rating_list)-1].save()
        except:
            print('error last')

        count_places = 1
        kafedra_rating_list = Rating.objects.filter(
            year=year, profile__kafedra=kafedra).order_by("-summ_effect")
        for i in range(len(kafedra_rating_list)-1):
            if kafedra_rating_list[i].summ_effect != kafedra_rating_list[i+1].summ_effect:
                kafedra_rating_list[i].kafedraplace_effect = count_places
                count_places += 1
            else:
                kafedra_rating_list[i].kafedraplace_effect = count_places
            kafedra_rating_list[i].save()

        try:
            if kafedra_rating_list[len(kafedra_rating_list)-1].summ_effect != kafedra_rating_list[len(kafedra_rating_list)-2].summ_effect:
                kafedra_rating_list[len(
                    kafedra_rating_list)-1].kafedraplace_effect = count_places
            else:
                kafedra_rating_list[len(
                    kafedra_rating_list)-1].kafedraplace_effect = count_places - 1
            kafedra_rating_list[len(kafedra_rating_list)-1].save()
        except:
            print('error last')

    # рейтинг по должностям всех сотрудников
    for dolzhnost in dolzhnost_list:
        print("idem po" + dolzhnost)
        count_places = 1
        dolzhnost_rating_list = Rating.objects.filter(year=year, profile__info__dolznost=dolzhnost).order_by(
            "-summ_effect").exclude(effect_stavka__lte=0.8)
        for i in range(len(dolzhnost_rating_list)-1):
            if dolzhnost_rating_list[i].summ_effect != dolzhnost_rating_list[i+1].summ_effect:
                dolzhnost_rating_list[i].dolzhnostplace = count_places
                count_places += 1
            else:
                dolzhnost_rating_list[i].dolzhnostplace = count_places
            print(dolzhnost_rating_list[i].profile.info.dolznost)
            print(dolzhnost_rating_list[i].profile.fullname)
            print(dolzhnost_rating_list[i].summ)
            print(dolzhnost_rating_list[i].dolzhnostplace)
            dolzhnost_rating_list[i].save()
        try:
            if dolzhnost_rating_list[len(dolzhnost_rating_list)-1].summ_effect != dolzhnost_rating_list[len(dolzhnost_rating_list)-2].summ_effect:
                dolzhnost_rating_list[len(
                    dolzhnost_rating_list)-1].dolzhnostplace = count_places
            else:
                dolzhnost_rating_list[len(
                    dolzhnost_rating_list)-1].dolzhnostplace = count_places - 1
            dolzhnost_rating_list[len(dolzhnost_rating_list)-1].save()
        except:
            print('error last')

        count_places = 1
        dolzhnost_rating_list = Rating.objects.filter(
            year=year, profile__info__dolznost=dolzhnost).order_by("-summ_effect")
        for i in range(len(dolzhnost_rating_list)-1):
            if dolzhnost_rating_list[i].summ_effect != dolzhnost_rating_list[i+1].summ_effect:
                dolzhnost_rating_list[i].dolzhnostplace_effect = count_places
                count_places += 1
            else:
                dolzhnost_rating_list[i].dolzhnostplace_effect = count_places
            dolzhnost_rating_list[i].save()
        try:
            if dolzhnost_rating_list[len(dolzhnost_rating_list)-1].summ_effect != dolzhnost_rating_list[len(dolzhnost_rating_list)-2].summ_effect:
                dolzhnost_rating_list[len(
                    dolzhnost_rating_list)-1].dolzhnostplace_effect = count_places
            else:
                dolzhnost_rating_list[len(
                    dolzhnost_rating_list)-1].dolzhnostplace_effect = count_places - 1
            dolzhnost_rating_list[len(dolzhnost_rating_list)-1].save()
        except:
            print('error last')


"""Функции для заполнения бд из нагрузок"""


def zapolnenie_bd():
    print('Введите название папки в с файлами в plan/, ./dir например (paht to dir with files) или название кафедры н-р администр.права')
    path_input = str(input())
    try:
        kafedra = Kafedra.objects.get(fullname=path_input)
        print('заполняем нагрузку из файла на сервере')
        print('Введите тип нагрузки 0=планируемая,1=фактическая(type of nagtuzka)')
        type = int(input())
        print('Введите год нагрузки, 2019 например (year of nagtuzka)')
        year = int(input())
        print('Введите 0 если полностью перезаливаем, 1 если доливаем тому у кого нет')
        zalivka = int(input())exi
        profiles = Profile.objects.filter(kafedra=kafedra)
        data = None
        logs = open("logs.txt", 'a')
        status = None
        data = ""
        if type == 0:
            status = False
        else:
            status = True
        for p in profiles:
            if Predmet.objects.filter(prepodavatel=p, year=int(year), status=status).count() > 0 and zalivka == 1:
                    print('пропускаем ' + p.fullname+' уже есть предметы'+str(
                        Predmet.objects.filter(prepodavatel=p, year=int(year), status=status).count()))
                    continue
            try:
                if type == 0:
                    nagruzkadoc = Nagruzka.objects.get(year=year, kafedra=kafedra,status='Планируемая')
                    data = takeXls(nagruzkadoc.document.path, ''.join([p.fullname.split(' ')[0], ' ', p.fullname.split(' ')[1][0], '.',
                                          p.fullname.split(' ')[2][0]]), True) 
                else:
                    nagruzkadoc = Nagruzka.objects.get(year=year, kafedra=kafedra,status='Фактическая')
                    data = takeXls(nagruzkadoc.document.path, ''.join([p.fullname.split(' ')[0], ' ', p.fullname.split(' ')[1][0], '.',
                                          p.fullname.split(' ')[2][0]]), False) 
                if not isinstance(data, list):
                    try:
                        print(p.kafedra.fullname+"  "+p.fullname)
                        print(data)
                        logs.write(p.kafedra.fullname+"  " +
                                    p.fullname + " " + data+'\n')
                        continue
                    except Exception as e:
                        logs.write(str(e)+" \n")
                try:
                    fields = Predmet._meta.get_fields()
                    for table in range(len(data)):
                        if table == 0:
                            for row in range(len(data[table])):
                                count = 0
                                predmet = Predmet()
                                # print(data)
                                # print(len(data[table]))
                                # print(data[table])

                                if data[table][row][0] == "0":
                                    continue
                                for field in fields:
                                    if field.name == "id":
                                        continue
                                    if row == (len(data[table])-1) and field.name == "name":
                                        setattr(predmet, field.name,
                                                "Итого за 1 полугодие:")
                                        continue
                                    if row == (len(data[table])-1) and field.name == "auditor_nagruzka":
                                        setattr(predmet, field.name,
                                                data[table][row][count])
                                        break
                                    # print("suka")
                                    setattr(predmet, field.name,
                                            data[table][row][count])

                                    if count == 25:
                                        break
                                    count += 1
                                predmet.kafedra = p.kafedra
                                # print(predmet.__dict__)
                                predmet.prepodavatel = p
                                predmet.year = year
                                predmet.polugodie = '1'
                                if type == 0:
                                    predmet.status = False
                                else:
                                    predmet.status = True
                                # tru если выполена false если план
                                # print(data)
                                predmet.save1()

                        if table == 1:
                            for row in range(len(data[table])):
                                count = 0
                                predmet = Predmet()
                                # print(len(data[table]))
                                # print(data[table])
                                # print( data[table][row][0])
                                if data[table][row][0] == "0":

                                    continue
                                for field in fields:
                                    if field.name == "id":
                                        continue
                                    if row == (len(data[table])-2) and field.name == "name":
                                        setattr(predmet, field.name,
                                                "Итого за 2 полугодие:")
                                        continue
                                    if row == (len(data[table])-1) and field.name == "name":
                                        setattr(predmet, field.name,
                                                "Итого за учебный год:")
                                        continue

                                    if row == (len(data[table])-2) and field.name == "auditor_nagruzka":
                                        setattr(predmet, field.name,
                                                data[table][row][count])
                                        break
                                    if row == (len(data[table])-1) and field.name == "auditor_nagruzka":
                                        setattr(predmet, field.name,
                                                data[table][row][count])
                                        break
                                    # print("suka2")
                                    setattr(predmet, field.name,
                                            data[table][row][count])
                                    # print(count)
                                    if count == 25:
                                        break
                                    count += 1

                                predmet.kafedra = p.kafedra
                                # print(predmet.__dict__)
                                predmet.prepodavatel = p
                                predmet.year = year
                                predmet.polugodie = '2'
                                if type == 0:
                                    predmet.status = False
                                else:
                                    predmet.status = True
                                # tru если выполена false если план
                                # print(predmet.all_values())
                                predmet.save1()

                    print("успешно обработан план "+p.fullname)
                except Exception as e:
                    logs.write("неудалось в профиле " +
                            p.fullname + " "+str(e)+'\n')

                    continue

            except Exception as e:

                logs.write(kafedra.fullname +
                        " не прошел документ " + str(e)+'\n')
                logs.write(full_path + '\\' + file_name + " " +
                        p.fullname.split(' ', 1)[0] + '\n')
                continue
    except Exception as e:
        print(e)
        path = "./plan/" + path_input
        file_list = os.listdir(path)
        print(file_list)
        print('Введите тип нагрузки 0=планируемая,1=фактическая(type of nagtuzka)')
        type = int(input())
        print('Введите год нагрузки, 2019 например (year of nagtuzka)')
        year = int(input())
        data = None
        for file_name in file_list:
            print(file_name)
            kafedra = None
            try:
                if "АП" in file_name:
                    kafedra = Kafedra.objects.get(fullname="администр.права")
                if "АД" in file_name:
                    kafedra = Kafedra.objects.get(fullname="администрат.деят.ОВД")
                if "ГиТП" in file_name:
                    kafedra = Kafedra.objects.get(
                        fullname="граж.и труд.права,гражд.процесса")
                if "ДОУ" in file_name:
                    kafedra = Kafedra.objects.get(fullname="деят. ОВД в ОУ")
                if "ИД" in file_name:
                    kafedra = Kafedra.objects.get(fullname="исслед. докум.")
                if "Ин яз" in file_name:
                    kafedra = Kafedra.objects.get(fullname="иностр. языков")
                if "ИиМ" in file_name:
                    kafedra = Kafedra.objects.get(fullname="инф. и мат.")
                if "ИБ" in file_name:
                    kafedra = Kafedra.objects.get(fullname="информац. без-ти")
                if "ИГиП" in file_name:
                    kafedra = Kafedra.objects.get(fullname="истории гос.и пр.")
                if "КиМП" in file_name:
                    kafedra = Kafedra.objects.get(fullname="конст.и муниц.пр.")
                if "Криминалистика" in file_name:
                    kafedra = Kafedra.objects.get(fullname="криминалистики")
                if "Криминология" in file_name:
                    kafedra = Kafedra.objects.get(fullname="криминологии")
                if "ОП" in file_name:
                    kafedra = Kafedra.objects.get(fullname="огневой подгот.")
                if "ОРД" in file_name:
                    kafedra = Kafedra.objects.get(fullname="ОРД и спец.техники")
                if "ОиТ" in file_name:
                    kafedra = Kafedra.objects.get(
                        fullname="оружиевед.и трасологии")
                if "Педагогики" in file_name or "Педагогика" in file_name:
                    kafedra = Kafedra.objects.get(fullname="педагогики")
                if "ПЧиМП" in file_name:
                    kafedra = Kafedra.objects.get(
                        fullname="прав человека и междун.права")
                if "ПР" in file_name:
                    kafedra = Kafedra.objects.get(fullname="предв. расслед.")
                if "психология" in file_name:
                    kafedra = Kafedra.objects.get(fullname="психологии")
                if "Русск.яз" in file_name:
                    kafedra = Kafedra.objects.get(fullname="русского языка")
                if "СиП" in file_name:
                    kafedra = Kafedra.objects.get(fullname="социол.и полит.")
                if "СТ" in file_name:
                    kafedra = Kafedra.objects.get(fullname="специальной тактики")
                if "СИТ" in file_name:
                    kafedra = Kafedra.objects.get(fullname="спец.инф.технол.")
                if "ТГП" in file_name:
                    kafedra = Kafedra.objects.get(fullname="теории гос.и права")
                if "ТКОЭИ" in file_name:
                    kafedra = Kafedra.objects.get(fullname="ТКОЭИ")
                if "УП" in file_name:
                    kafedra = Kafedra.objects.get(fullname="уголовного права")
                if "Упр" in file_name:
                    kafedra = Kafedra.objects.get(fullname="уголовного процесса")
                if "ФП" in file_name:
                    kafedra = Kafedra.objects.get(fullname="физ. подготовки")
                if "Философии" in file_name:
                    kafedra = Kafedra.objects.get(fullname="философии")
                if "ЭиБУ" in file_name:
                    kafedra = Kafedra.objects.get(fullname="экономики и бух.учета")
                if "КЭБФиЭА" in file_name:
                    kafedra = Kafedra.objects.get(
                        fullname="эконом.безоп.,финансов и эконом.анализа")
                if "ЭКД" in file_name:
                    kafedra = Kafedra.objects.get(fullname="эксп.крим.деят-ти")
                if "Юр псих" in file_name:
                    kafedra = Kafedra.objects.get(
                        fullname="юридической психологии")
            except Exception as e:
                print(file_name)

            full_path = os.path.abspath(os.curdir)+'/plan/'+path_input
            # print(full_path)
            profiles = Profile.objects.filter(kafedra=kafedra)
            logs = open("logs.txt", 'a')
            status = None
            data = ""
            if type == 0:
                status = False
            else:
                status = True
            for p in profiles:
                if Predmet.objects.filter(prepodavatel=p, year=int(year), status=status).count() > 0:
                    print('пропускаем ' + p.fullname+' уже есть предметы'+str(
                        Predmet.objects.filter(prepodavatel=p, year=int(year), status=status).count()))
                    continue

                try:
                    if type == 0:
                        data = takeXls(full_path+'/'+file_name,
                                    p.fullname.split(' ', 1)[0], True)
                    else:
                        data = takeXls(full_path+'/'+file_name,
                                    p.fullname.split(' ', 1)[0], False)
                    if not isinstance(data, list):
                        try:
                            print(p.kafedra.fullname+"  "+p.fullname)
                            print(data)
                            logs.write(p.kafedra.fullname+"  " +
                                        p.fullname + " " + data+'\n')
                            continue
                        except Exception as e:
                            logs.write(str(e)+" \n")
                    try:
                        fields = Predmet._meta.get_fields()
                        for table in range(len(data)):
                            if table == 0:
                                for row in range(len(data[table])):
                                    count = 0
                                    predmet = Predmet()
                                    # print(data)
                                    # print(len(data[table]))
                                    # print(data[table])

                                    if data[table][row][0] == "0":
                                        continue
                                    for field in fields:
                                        if field.name == "id":
                                            continue
                                        if row == (len(data[table])-1) and field.name == "name":
                                            setattr(predmet, field.name,
                                                    "Итого за 1 полугодие:")
                                            continue
                                        if row == (len(data[table])-1) and field.name == "auditor_nagruzka":
                                            setattr(predmet, field.name,
                                                    data[table][row][count])
                                            break
                                        # print("suka")
                                        setattr(predmet, field.name,
                                                data[table][row][count])

                                        if count == 25:
                                            break
                                        count += 1
                                    predmet.kafedra = p.kafedra
                                    # print(predmet.__dict__)
                                    predmet.prepodavatel = p
                                    predmet.year = year
                                    predmet.polugodie = '1'
                                    if type == 0:
                                        predmet.status = False
                                    else:
                                        predmet.status = True
                                    # tru если выполена false если план
                                    # print(data)
                                    predmet.save1()

                            if table == 1:
                                for row in range(len(data[table])):
                                    count = 0
                                    predmet = Predmet()
                                    # print(len(data[table]))
                                    # print(data[table])
                                    # print( data[table][row][0])
                                    if data[table][row][0] == "0":

                                        continue
                                    for field in fields:
                                        if field.name == "id":
                                            continue
                                        if row == (len(data[table])-2) and field.name == "name":
                                            setattr(predmet, field.name,
                                                    "Итого за 2 полугодие:")
                                            continue
                                        if row == (len(data[table])-1) and field.name == "name":
                                            setattr(predmet, field.name,
                                                    "Итого за учебный год:")
                                            continue

                                        if row == (len(data[table])-2) and field.name == "auditor_nagruzka":
                                            setattr(predmet, field.name,
                                                    data[table][row][count])
                                            break
                                        if row == (len(data[table])-1) and field.name == "auditor_nagruzka":
                                            setattr(predmet, field.name,
                                                    data[table][row][count])
                                            break
                                        # print("suka2")
                                        setattr(predmet, field.name,
                                                data[table][row][count])
                                        # print(count)
                                        if count == 25:
                                            break
                                        count += 1

                                    predmet.kafedra = p.kafedra
                                    # print(predmet.__dict__)
                                    predmet.prepodavatel = p
                                    predmet.year = year
                                    predmet.polugodie = '2'
                                    if type == 0:
                                        predmet.status = False
                                    else:
                                        predmet.status = True
                                    # tru если выполена false если план
                                    # print(predmet.all_values())
                                    predmet.save1()

                        print("успешно обработан план "+p.fullname)
                    except Exception as e:
                        logs.write("неудалось в профиле " +
                                p.fullname + " "+str(e)+'\n')

                        continue

                except Exception as e:

                    logs.write(kafedra.fullname +
                            " не прошел документ " + str(e)+'\n')
                    logs.write(full_path + '\\' + file_name + " " +
                            p.fullname.split(' ', 1)[0] + '\n')
                    continue
            continue
    return 0


def create_plans():
    print("enter year for new plans")
    year = int(input())
    profiles = Profile.objects.all()
    count = 0
    for p in profiles:
        if not Plan.objects.get(prepod=p, year=year):
            newplan = Plan()
            newplan.user = p.user
            newplan.prepod = p
            newplan.year = year

            newplan.name = "".join([p.fullname.split(' ')[0], ' ', p.fullname.split(' ')[
                                   1][0], '.', p.fullname.split(' ')[2][0]])
            print("sozadanie plana ")
            count += 1
            # newplan.save()
    print("sozdano "+count+" planov iz " + profiles.count()+"profiley")


def utverzhdDocument():
    print("введдите 1 если для всех кафедр, 2 если для конкретной")
    commands = input()
    print("введдите год")
    year = int(input())
    if commands == "2":
        print("Введите название кафедры как на сайте")
        kafedra_name=input()
        kafedra = Kafedra.objects.get(fullname=kafedra_name)
        rating_list = Rating.objects.filter(profile__kafedra=kafedra, year=year)
        data = []
        kafedra_result =[0,0,0,0,0]
        for rating in rating_list:
            fio =''.join(
            [rating.profile.fullname.split(' ')[0], ' ', rating.profile.fullname.split(' ')[1][0], '.',
             rating.profile.fullname.split(' ')[2][0]])
            buff =[
            fio,
            rating.profile.info.dolznost,
            rating.urr,
            rating.ormr, 
            rating.mrr,
            rating.pcr,
            rating.summ, 
            rating.kafedraplace, 
            rating.dolzhnostplace,
            rating.unikplace
             ]
            data.append(buff)
            kafedra_result[0] += rating.urr
            kafedra_result[1] += rating.ormr
            kafedra_result[2] += rating.mrr
            kafedra_result[3] += rating.pcr
            kafedra_result[4] += rating.summ
        print(kafedra.fullname,data,kafedra_result)
        createDocument(kafedra.fullname,data,kafedra_result)
             

        
    elif commands == '1':
        kafedra_list= Kafedra.objects.all()
        for kafedra in kafedra_list:
            rating_list = Rating.objects.filter(profile__kafedra=kafedra, year=year)
            data = []
            kafedra_result =[0,0,0,0,0]
            for rating in rating_list:
                try:
                    fio =''.join(
                    [rating.profile.fullname.split(' ')[0], ' ', rating.profile.fullname.split(' ')[1][0], '.',
                    rating.profile.fullname.split(' ')[2][0]])
                except Exception as e:
                    print(e)
                    fio = rating.profile.fullname.split(' ')[0]
                try:
                    dolznost = rating.profile.info.dolznost
                except Exception as e:
                    print(e)
                    dolznost = "не назначено"
                buff =[
                fio,
                dolznost,
                float("{0:.2f}".format(rating.urr)),
                float("{0:.2f}".format(rating.ormr)),
                float("{0:.2f}".format(rating.mrr)),
                float("{0:.2f}".format(rating.pcr)),
                float("{0:.2f}".format(rating.summ)), 
                rating.kafedraplace_effect, 
                rating.dolzhnostplace_effect,
                rating.unikplace_effect,
                ]
                data.append(buff)
                kafedra_result[0] += rating.urr
                kafedra_result[1] += rating.ormr
                kafedra_result[2] += rating.mrr
                kafedra_result[3] += rating.pcr
                kafedra_result[4] += rating.summ
            createDocument(kafedra.fullname,data,kafedra_result)


def kafedraRatingList(year):
    kafedras = Kafedra.objects.all()
    graphdata = []
    graphdata_kafedra_list = []
    buff = []
    urr_list = []
    ormr_list = []
    mrr_list = []
    pcr_list = []
    summ_list = []
    kafedra_list = []
    urr = 0
    ormr = 0
    mrr = 0
    pcr = 0
    kafedra_summ = 0
    effect_stavka_summ = 0
    for k in kafedras:
        profiles = k.prepods.all()
        for p in profiles:
            try:
                rating = Rating.objects.get(profile=p, year=year)
                urr += rating.urr
                ormr += rating.ormr
                mrr += rating.mrr
                pcr += rating.pcr
                kafedra_summ += rating.urr+rating.ormr+rating.mrr+rating.pcr
                effect_stavka_summ += rating.effect_stavka
            except Rating.DoesNotExist:
                continue
            except Exception as e:
                print(e)
                continue
        try:
            summ_list.append(
                (float("{0:.2f}".format(kafedra_summ/effect_stavka_summ))))
            urr_list.append(
                (float("{0:.2f}".format(urr/effect_stavka_summ))))
            ormr_list.append(
                (float("{0:.2f}".format(ormr/effect_stavka_summ))))
            mrr_list.append(
                (float("{0:.2f}".format(mrr/effect_stavka_summ))))
            pcr_list.append(
                (float("{0:.2f}".format(pcr/effect_stavka_summ))))
            kafedra_list.append(k.fullname)
        except Exception as e:
            print(e)

        urr = 0
        ormr = 0
        mrr = 0
        pcr = 0
        kafedra_summ = 0
        effect_stavka_summ = 0
    for i in range(len(summ_list)-1):
        for j in range(len(summ_list)-2, i-1, -1):
            if summ_list[j+1] > summ_list[j]:
                summ_list[j], summ_list[j +
                                        1] = summ_list[j+1], summ_list[j]
                urr_list[j], urr_list[j +
                                        1] = urr_list[j+1], urr_list[j]
                ormr_list[j], ormr_list[j +
                                        1] = ormr_list[j+1], ormr_list[j]
                mrr_list[j], mrr_list[j +
                                        1] = mrr_list[j+1], mrr_list[j]
                pcr_list[j], pcr_list[j +
                                        1] = pcr_list[j+1], pcr_list[j]
                kafedra_list[j], kafedra_list[j +
                                                1] = kafedra_list[j+1], kafedra_list[j]
    graphdata_kafedra_list.append({
        "name": 'Кафедры',
        "data": kafedra_list
    })
    graphdata.append({
        "name": 'Учебная работа',
        "data": urr_list
    })
    graphdata.append({
        "name": 'Организационно методическая работа',
        "data": ormr_list
    })
    graphdata.append({
        "name": 'Подготовка учебно-методических материалов',
        "data": mrr_list
    })
    buff.append(pcr)
    graphdata.append({
        "name": 'Педагогический контроль',
        "data": pcr_list
    })
    response = []
    count = 0
    for kafedra in kafedra_list:
        response.append([kafedra,urr_list[count],ormr_list[count],mrr_list[count],pcr_list[count],summ_list[count],count+1])
        count += 1
    
    print(response)
    writeAllKafDoc(year,response)

    


print("Выберете опцию: 1 запонить нагрузки из папки , 2 создать планы ,3 обновить рейтинги",
      "4 порядок месяцев", "5 сформировать файлы утверждения рейтингов","6 рейтинг кафедр док")
command=input()
if command == '1':
    zapolnenie_bd()
elif command == '2':
    create_plans()
elif command == '3':
    print("введите год")
    year=input()
    refresh_places(int(year))
elif command == '4':
    mess_order()
elif command == '5':
    utverzhdDocument()
elif command == '6':
    print("введите год")
    year=input()
    kafedraRatingList(int(year))

print('konets')

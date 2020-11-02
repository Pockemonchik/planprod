# -*- coding: utf8 -*-
'''
Установка дополнительных библиотек: pip install python-docx, pip install openxls, pip install PyPDF2
Версия библиотеки: 0.8.10
Зависимости: Python 2.6, 2.7, 3.3, or 3.4     lxml >= 2.3.2
Документация: https://python-docx.readthedocs.io/en/latest/
'''
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from django.utils.lorem_ipsum import paragraph, paragraphs
from docx.shared import Pt
import re
import openpyxl
import sys
import os
import subprocess
# import comtypes.client
# from PyPDF2 import PdfFileReader


# Функция для конвертации *.docx в *.pdf и возврат результата
def convert(nameDoc):
    wdFormatPDF = 17

    in_file = os.path.abspath('C:/Users/UzziMauzer/Desktop/planprod/mvd/plan' + nameDoc)
    out_file = os.path.abspath('C:/Users/UzziMauzer/Desktop/planprod/mvd/plan' + "out.pdf")

    word = comtypes.client.CreateObject('Word.Application')
    doc = word.Documents.Open(in_file)
    doc.SaveAs(out_file, FileFormat=wdFormatPDF)
    doc.Close()
    word.Quit()

    pdf_document = "out.pdf"
    pdf = PdfFileReader(open('C:/Users/UzziMauzer/Desktop/planprod/mvd/plan' + pdf_document, "rb"))

    return pdf

#Функция для чтения списка преподавателей
def xlsPrepod(nameDoc):
    wb = openpyxl.load_workbook(filename =nameDoc, data_only=True)

    sheets = wb.sheetnames
    ws = wb[sheets[0]]
    listAll = []
    numString = 4
    numIndex = 0

    while(numString<771):
        if str(ws['C'+str(numString)].value) == 'None':
            numString+=1
        else:
            listAll.append([])
            listAll[numIndex].append(str(ws['C'+str(numString)].value))
            listAll[numIndex].append(str(ws['D'+str(numString)].value))
            listAll[numIndex].append(str(ws['E'+str(numString)].value) + ' ' + str(ws['F'+str(numString)].value) + ' ' + str(ws['G'+str(numString)].value))
            listAll[numIndex].append(str(ws['K'+str(numString)].value))
            numIndex+=1
            numString += 1

    return listAll

# Функция для получения упорядоченного списка всех нужных переменных из xls файла и его возврат
def takeXls(nameDoc,namePrepod,flag):
    #wb = openpyxl.load_workbook(filename =  '/home/andrey/Documents/Vazhno_sho_pizdec/mvdproject/mvd/plan/' + nameDoc + '.xlsx', data_only=True)
    wb = openpyxl.load_workbook(filename = nameDoc , data_only=True)
    '''
    needColumnFirst = ['O','Q','S','U','W','X','Y','AA','AC','AE','AG','AI','AK','AM','AO','AQ','AS','AU','AW','AY','BA','BC','BE','BF','BG' ]
    needColumnSecond = ['CA','CC','CE','CG','CI','CJ','CK','CM','CO','CQ','CS','CU','CW','CY','DA','DC','DE','DG','DI','DK','DM','DO','DQ','DS','DT']
    needColumnThird = ['EM','EO','EQ','ES','EU','EV','EW','EY','FA','FC','FE','FG','FI','FK','FM','FO','FQ','FS','FU','FW','FY','GA','GC','GE','GF']
    '''
    if flag==True:
        zapasFact=[]
        zapasFactAudit=[]
        nameNeedColunm = ['лекции', 'семинары', 'практические занятия в группе', 'практические занятия в подгруппе', 'круглый стол', 'консультации перед экзаменами', 'текущие консультации', 'внеаудиторное чтение', 'практика руководство', 'ВКР', 'курсовая работа', 'контрольная работа аудиторная', 'контрольная работа домашняя', 'проверка практикума', 'проверка лабораторной работы', 'защита практики', 'зачет устный', 'зачет письменный', 'вступительные испытания', 'экзамены', 'государственные экзамены', 'юнктура)', 'руководство адъюнктами', 'Плановая нагрузка', 'Плановая аудиторная нагрузка']
    else:
        nameNeedColunm = ['лекции', 'семинары', 'практические занятия в группе', 'практические занятия в подгруппе', 'круглый стол', 'консультации перед экзаменами', 'текущие консультации', 'внеаудиторное чтение', 'практика руководство', 'ВКР', 'курсовая работа', 'контрольная работа аудиторная', 'контрольная работа домашняя', 'проверка практикума', 'проверка лабораторной работы', 'защита практики', 'зачет устный', 'зачет письменный', 'вступительные испытания', 'экзамены', 'государственные экзамены', 'юнктура)', 'руководство адъюнктами', 'Фактически выполненная нагрузка', 'Фактически выполненная аудиторная нагрузка']
        zapasFact=['Фактически выполненная нагрузка', 'выполненная нагрузка','выполненая нагрузка','Фактическая нагрузка']
        zapasFactAudit=['Фактически выполненная аудиторная нагрузка', 'выполненная аудиторная нагрузка','выполненая аудиторная нагрузка','Фактическая аудиторная нагрузка']


    nameprepodColumn = ['наименование дисциплины', 'Факультет', 'Курс, группа']



    listFirst = []
    listSecond = []
    listResult = []

    needColumnFirst = []
    needColumnSecond = []
    needColumnThird = []
    needPrepodColumn = []


    sheets = wb.sheetnames
    ws = wb[sheets[0]]
    prepodRow = 1

    #Защита от дураков на ошибки
    for textCol in zapasFact:
        for column in ws.iter_cols(min_col=1, min_row=8, max_col=500, max_row=8):
            if textCol.lower() in str(column[0].value).lower():
                print(column[0].value)
                column[0].value = 'фактически выполненная нагрузка'
                print(column[0].value)

    for textCol in zapasFactAudit:
        for column in ws.iter_cols(min_col=1, min_row=8, max_col=500, max_row=8):
                if textCol.lower() in str(column[0].value).lower():
                    column[0].value = 'фактически выполненная аудиторная нагрузка'

    

    #Функция заполнения Индексов столбцов
    for textCol in nameNeedColunm:
        k = 0
        for column in ws.iter_cols(min_col=1, min_row=8, max_col=200, max_row=8):
                if textCol.lower() in str(column[0].value).lower():
                    if textCol.lower() == 'экзамены':
                        if textCol.lower() != str(column[0].value).lower():
                            continue
                    k += 1
                    column[0].value = ' '
                    reg = re.search(".+?\.(.+?)8>", str(column[0]))
                    if reg and len(reg.group(1)) > 0:
                        needText = reg.group(1)
                    if k == 1:
                        needColumnFirst.append(needText)
                    elif k == 2:
                        needColumnSecond.append(needText)

                    else: needColumnThird.append(needText)
        if k != 3 :
            print(k)
            print(textCol)
            print('error xls')
            return textCol

    

    for textCol in nameprepodColumn:
        k = 0
        for column in ws.iter_cols(min_col=1, min_row=8, max_col=200, max_row=8):
                if textCol.lower() == str(column[0].value).lower():
                    k += 1
                    column[0].value = ' '
                    reg = re.search(".+?\.(.+?)8>", str(column[0]))
                    if reg and len(reg.group(1)) > 0:
                        needText = reg.group(1)
                    needPrepodColumn.append(needText)




    if len(needColumnFirst) != len(nameNeedColunm) or len(needColumnSecond) != len(nameNeedColunm) or len(needColumnThird) != len(nameNeedColunm):
        print(len(needColumnFirst) , ' ' , len(nameNeedColunm) , ' ' , len(needColumnSecond) , ' '  , len(nameNeedColunm) , ' ' , len(needColumnThird) , ' ' , len(nameNeedColunm))
        print('error xls file')
        return("error xls file")




    #Находим строку с преподом
    print(str(namePrepod).lower())
    for column in ws.iter_rows(min_col=2, min_row=1, max_col=2, max_row=1000, values_only=True):
        if str(namePrepod).lower() in str(column[0]).lower():
            break
        prepodRow += 1
    
    if prepodRow == 1001:
        return ("404")

    secondPrepod = 1
    count = 0

    #Находим строку с следующим преподом
    for column in ws.iter_rows(min_col=2, min_row=(prepodRow+1), max_col=2, max_row=(prepodRow+30), values_only=True):
        reg = re.search("([А-ЯЁ].+?\s)[А-ЯЁ]\.[А-ЯЁ]\.", str(column[0]))
        if reg and len(reg.group(1)) > 0:
            count = secondPrepod
            break
        reg = re.search("(ИТОГО):", str(column[0]))
        if reg and len(reg.group(1)) > 0:
            count = secondPrepod
            break

        secondPrepod+=1

    if count == 0:
        count = 14



    #Получаем значения 1 таблицы в лист


    for takeStr in range(1,count):
        listFirst.append([])
        if str(ws[str(needPrepodColumn[0])+str(prepodRow+takeStr)].value) != 'None':
            firstBigSell = str(ws[str(needPrepodColumn[0])+str(prepodRow+takeStr)].value)
            if str(ws[str(needPrepodColumn[3])+str(prepodRow+takeStr)].value) != 'None':
                firstBigSell += ' ' + str(ws[str(needPrepodColumn[3])+str(prepodRow+takeStr)].value)
                if str(ws[str(needPrepodColumn[6])+str(prepodRow+takeStr)].value) != 'None':
                    firstBigSell += ' ' + str(ws[str(needPrepodColumn[6])+str(prepodRow+takeStr)].value)
            else:
                if str(ws[str(needPrepodColumn[6])+str(prepodRow+takeStr)].value) != 'None':
                    firstBigSell += ' ' + str(ws[str(needPrepodColumn[6])+str(prepodRow+takeStr)].value)
            listFirst[takeStr-1].append(firstBigSell)
        else:
            listFirst[takeStr-1].append('0')

        for takeColumn in needColumnFirst:
            cell = takeColumn+str(prepodRow+takeStr)
            if ws[cell].value == None or ws[cell].value == '0':
                listFirst[takeStr-1].append('0')
            else:
                listFirst[takeStr-1].append(ws[cell].value)
    listFirst.append([])
    for takeColumn in needColumnFirst:
            cell = takeColumn+str(prepodRow)
            if ws[cell].value == None or ws[cell].value == '0':
                listFirst[takeStr].append('0')
            else:
                listFirst[takeStr].append(ws[cell].value)

    #Получаем значения 2 таблицы в лист
    for takeStr in range(1,count):
        listSecond.append([])




        if str(ws[str(needPrepodColumn[1])+str(prepodRow+takeStr)].value) != 'None':
            firstBigSell = str(ws[str(needPrepodColumn[1])+str(prepodRow+takeStr)].value)
            if str(ws[str(needPrepodColumn[4])+str(prepodRow+takeStr)].value) != 'None':
                firstBigSell += ' ' + str(ws[str(needPrepodColumn[4])+str(prepodRow+takeStr)].value)
                if str(ws[str(needPrepodColumn[7])+str(prepodRow+takeStr)].value) != 'None':
                    firstBigSell += ' ' + str(ws[str(needPrepodColumn[7])+str(prepodRow+takeStr)].value)
            else:
                if str(ws[str(needPrepodColumn[7])+str(prepodRow+takeStr)].value) != 'None':
                    firstBigSell += ' ' + str(ws[str(needPrepodColumn[7])+str(prepodRow+takeStr)].value)
            listSecond[takeStr-1].append(firstBigSell)

        else:
            listSecond[takeStr-1].append('0')

        for takeColumn in needColumnSecond:
            cell = takeColumn+str(prepodRow+takeStr)
            if ws[cell].value == None or ws[cell].value == '0':
                listSecond[takeStr-1].append('0')
            else:
                listSecond[takeStr-1].append(ws[cell].value)
    listSecond.append([])
    for takeColumn in needColumnSecond:
            cell = takeColumn+str(prepodRow)
            if ws[cell].value == None or ws[cell].value == '0':
                listSecond[takeStr].append('0')
            else:
                listSecond[takeStr].append(ws[cell].value)
    listSecond.append([])
    #Получаем значения 3 таблицы в лист
    for takeColumn in needColumnThird:
            cell = takeColumn+str(prepodRow)
            if ws[cell].value == None or ws[cell].value == '0':
                listSecond[takeStr+1].append('0')
            else:
                listSecond[takeStr+1].append(ws[cell].value)

    listResult.append(listFirst)
    listResult.append(listSecond)

    for i in range(len(listResult)):
        for j in range(len(listResult[i])):
            print(listResult[i][j])

    return(listResult)

# Запись в документ Основной информации
def writeInfoDoc(listInfo, listCell,indexRow):
    # document = Document("C:/Users/UzziMauzer/Desktop/planprod/mvd/plan/docDynamic.docx")
    #
    document = Document("C:/Users/Andrey/Desktop/planprod/mvd/plan/docDynamic.docx")
    #document = Document("//home/user/planprod/mvd/plan/docDynamic.docx")
    # document = Document("C:/Users/UzziMauzer/Desktop/planprod/mvd/plan/docDynamic.docx")

    needStrings = [8, 9, 10, 20, 22, 23,31]
    # Для строчек, который разделяются на 3, то есть кафедры
    threeStrings = [25,27,29]
    threeList = listInfo.pop(6)
    oneOfThreeList = [[],[],[]]
    numberSimbolAllStrings = 0


    countParagraphs = 0
    countListInfo = 0
    for paragraphs in document.paragraphs:
        if countParagraphs in needStrings:
            if countListInfo != 6:
                paragraphs.text = (listInfo[countListInfo])
            else: paragraphs.text = '            '+(listInfo[countListInfo])+'          _'
            for run in paragraphs.runs:
                font = run.font
                font.size = Pt(14)
            if countListInfo == 6:
                font = run.font
                font.underline = True
                font.size = Pt(14)
            countListInfo+=1
            # if countParagraphs == 2 or countParagraphs == 3:
            #     for run in paragraphs.runs:
            #         font = run.font
            #         font.bold = 1



        # Для тех самых 3х строчек
        elif countParagraphs in threeStrings:
            for countStrings in range(0,3):
                if threeList != "":
                    if (len(threeList) - numberSimbolAllStrings) > 75:
                        flag = True
                        for indexSpase in range(75,0,-1):
                            if flag == True:
                                if threeList[indexSpase] == ' ':
                                    firstSpase = indexSpase
                                    numSimbol = 0
                                    while (numSimbol != firstSpase):
                                        oneOfThreeList[countStrings] += threeList[numSimbol + numberSimbolAllStrings]
                                        numSimbol +=1
                                    numberSimbolAllStrings += firstSpase + 1
                                    print(oneOfThreeList[countStrings])
                                    document.paragraphs[int(threeStrings[0])+countStrings*2].text = (oneOfThreeList[countStrings])
                                    for run in document.paragraphs[int(threeStrings[0])+countStrings*2].runs:
                                        font = run.font
                                        font.size = Pt(14)
                                    flag = False
                    else:
                        oneOfThreeList[countStrings] = threeList[numberSimbolAllStrings:]
                        document.paragraphs[threeStrings[countStrings]].text = (oneOfThreeList[countStrings])
                        for run in document.paragraphs[int(threeStrings[0])+countStrings*2].runs:
                            font = run.font
                            font.size = Pt(14)
                        threeList = ""
        countParagraphs += 1



    # document.save('C:/Users/UzziMauzer/Desktop/planprod/mvd/tmp.docx')


    #document.save("//home/user/planprod/mvd/plan/tmp.docx")


    #document.save("//root/planprod/mvd/plan/tmp.docx")
    document.save("C:/Users/Andrey/Desktop/planprod/mvd/plan/tmp.docx")
    #document.save("C:/Users/UzziMauzer/Desktop/planprod/mvd/plan/tmp.docx")
    return createDoc(listCell,indexRow)




# Чтение из документа основной информации
def readInfoDoc(nameDoc):
    document = Document(nameDoc)
    needStrings = [5,8,9,21,23,25,27,29,31]
    infoDoc = []

    for num in needStrings:
        if num == 27 or num == 28:
            infoDoc[5] += document.paragraphs[num].text
        else:
            infoDoc.append(document.paragraphs[num].text)
    print(infoDoc)

def createDoc2(nameDoc, listCell, numString1, numString2):
    #document = Document('//home/user/planprod/mvd/plan/test.docx')
    document = Document("C:/Users/Andrey/Desktop/planprod/mvd/plan/test.docx")
    #document = Document("test.docx")
    count = 0
    for table in document.tables:
        if count == 0:
            count +=1
            for i in range(0,numString1):
                table.add_row()
        else:
            for i in range(0,numString2):
                table.add_row()

    count = 0
    for table in document.tables:
        for row in (table.rows):
            for cell in row.cells:
                if cell.text.lower() == '':
                    cell.text = str(listCell[count])
                    cell.alignment=WD_ALIGN_PARAGRAPH.CENTER
#                     cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    count+=1
#         print(count)

    document.save('//home/user/planprod/mvd/plan/'+nameDoc+'.docx')

# Заполнение документа по полученным данным и сохранение документа
def createDoc(listCell, indexRow):
    try:
        #     document = Document('/home/andrey/Documents/Vazhno_sho_pizdec/mvdproject/mvd/plan/Shablon.docx')

        needTables = [1, 2, 3, 4, 6, 7, 8, 9, 10, 11, 12, 13, 14, 15,19]
        # document = Document("C:/Users/UzziMauzer/Desktop/planprod/mvd/tmp.docx")
        document = Document("C:/Users/Andrey/Desktop/planprod/mvd/plan/tmp.docx")
        #document = Document("//home/user/planprod/mvd/plan/tmp.docx")
        #document = Document("//root/planprod/mvd/plan/tmp.docx")
        print(len(document.tables))
        count = -1
        for i in needTables:
            count += 1
            table = document.tables[i-1]
            for j in range(indexRow[count]):
                table.add_row()
        needTables.append(5)
        needTables.append(16)
        needTables.sort()
        print(listCell)
        print(indexRow)

        count = 0
        for i in needTables:
            table = document.tables[i-1]
            for row in (table.rows):
                for cell in row.cells:
                    if cell.text.lower() == '' or cell.text.lower() =='x':
                        print(cell.text)
                        print(listCell[count])
                        cell.text = str(listCell[count])
                        cell.alignment=WD_ALIGN_PARAGRAPH.CENTER
        #                     cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                        count+=1

        #document.save(nameDoc+'.docx')
        # document.save("C:/Users/UzziMauzer/Desktop/asdasdasdasdasdasd"+'docx')
        print("Wow")
        return(document)
    except Exception as e:
        print(e)
#    return convert(nameDoc)
# def createDoc(nameDoc, listCell):
#     document = Document('C:/Users/UzziMauzer/Desktop/planprod/mvd/plan322.docx')
#
#     count = 0
#     for table in document.tables:
#
#         for row in (table.rows):
#             for cell in row.cells:
#                 if cell.text.lower()== 'x':
#                     cell.text = str(listCell[count])
#                     cell.alignment=WD_ALIGN_PARAGRAPH.CENTER
# #                     cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
#                     count+=1
# #         print(count)
#
#     document.save('C:/Users/UzziMauzer/Desktop/planprod/mvd/'+nameDoc+'.docx')
#     # return convert(nameDoc)
#     return document

#Для созданных этим скриптом таблиц (сreateTable)
def tekeUpdateTable(document):

    # document=Document('C:/Users/UzziMauzer/Downloads/planqwer.docx')
    listAllTable = []
    listOneTable = []
    listRow = []

    #Ежемас Таблица
    table = document.tables[4]
    countRow = 0
    for row in (table.rows):
        listRow = []
        if countRow > 0:
            for cell in row.cells:
                if cell.text == '' or cell.text == ' ':
                    listRow.append('0')
                else:
                    cell.text = cell.text.replace(',','.')
                    listRow.append(cell.text)
            listOneTable.append(listRow)
        countRow+=1
    listAllTable.append(listOneTable)
    print(listOneTable)
    listOneTable = []

    #Метод 1 полугодие
    table = document.tables[5]
    countRow = 0
    for row in (table.rows):
        listRow = []
        if countRow > 0:
            for cell in row.cells:
                listRow.append(cell.text)
            listOneTable.append(listRow)
        countRow+=1
    listAllTable.append(listOneTable)
    listOneTable = []

    #Метод за 2 полугодие
    table = document.tables[6]
    countRow = 0
    for row in (table.rows):
        listRow = []
        if countRow > 0:
            for cell in row.cells:
                listRow.append(cell.text)
            listOneTable.append(listRow)
        countRow+=1
    listAllTable.append(listOneTable)
    listOneTable = []

    #Научн 1 полугодие
    table = document.tables[7]
    countRow = 0
    for row in (table.rows):
        listRow = []
        if countRow > 0:
            for cell in row.cells:
                listRow.append(cell.text)
            listOneTable.append(listRow)
        countRow+=1
    listAllTable.append(listOneTable)
    listOneTable = []

    #Научн за 2 полугодие
    table = document.tables[8]
    countRow = 0
    for row in (table.rows):
        listRow = []
        if countRow > 0:
            for cell in row.cells:
                listRow.append(cell.text)
            listOneTable.append(listRow)
        countRow+=1
    listAllTable.append(listOneTable)
    listOneTable = []

    #Восп 1 полугодие
    table = document.tables[9]
    countRow = 0
    for row in (table.rows):
        listRow = []
        if countRow > 1:
            for cell in row.cells:
                listRow.append(cell.text)
            listOneTable.append(listRow)
        countRow+=1
    listAllTable.append(listOneTable)
    listOneTable = []

    #Восп за 2 полугодие
    table = document.tables[10]
    countRow = 0
    for row in (table.rows):
        listRow = []
        if countRow > 1:
            for cell in row.cells:
                listRow.append(cell.text)
            listOneTable.append(listRow)
        countRow+=1
    listAllTable.append(listOneTable)
    listOneTable = []

    #Иностр за 1 полугодие
    table = document.tables[11]
    countRow = 0
    for row in (table.rows):
        listRow = []
        if countRow > 1:
            for cell in row.cells:
                listRow.append(cell.text)
            listOneTable.append(listRow)
        countRow+=1
    listAllTable.append(listOneTable)
    listOneTable = []

    #Иностр за 2 полугодие
    table = document.tables[12]
    countRow = 0
    for row in (table.rows):
        listRow = []
        if countRow > 1:
            for cell in row.cells:
                listRow.append(cell.text)
            listOneTable.append(listRow)
        countRow+=1
    listAllTable.append(listOneTable)
    listOneTable = []

    #Иные за 1 полугодие
    table = document.tables[13]
    countRow = 0
    for row in (table.rows):
        listRow = []
        if countRow > 1:
            for cell in row.cells:
                listRow.append(cell.text)
            listOneTable.append(listRow)
        countRow+=1
    listAllTable.append(listOneTable)
    listOneTable = []

    #Иные за 2 полугодие
    table = document.tables[14]
    countRow = 0
    for row in (table.rows):
        listRow = []
        if countRow > 1:
            for cell in row.cells:
                listRow.append(cell.text)
            listOneTable.append(listRow)
        countRow+=1
    listAllTable.append(listOneTable)

    return listAllTable



# Ð’ Ð´Ð°Ð½Ð½Ð¾Ð¹ Ñ„ÑƒÐ½ÐºÑ†Ð¸Ð¸ Ð¼Ñ‹ ÐžÐ±Ñ€Ð°Ð±Ð°Ñ‚Ñ‹Ð²Ð°ÐµÐ¼ Ð¿Ð¾Ð»ÑƒÑ‡Ð°ÐµÐ¼Ñ‹Ð¹ Ð´Ð¾ÐºÑƒÐ¼ÐµÐ½Ñ‚ Ð¸ Ð²Ð¾Ð·Ð²Ñ€Ð°Ñ‰Ð°ÐµÐ¼ Ð²ÑÐµ Ñ‚Ð°Ð±Ð»Ð¸Ñ†Ñ‹ Ð²Ð¾ Ð²Ð»Ð¾Ð¶ÐµÐ½Ð½Ñ‹Ñ… ÑÐ¿Ð¸ÑÐºÐ°Ñ…
# Ð¡Ð¾Ð²ÐµÑ‚ÑƒÑŽ Ñ‚ÑƒÑ‚ Ð½Ð¸Ñ‡ÐµÐ³Ð¾ Ð½Ðµ Ð¼ÐµÐ½ÑÑ‚ÑŒ, Ð³Ð¾Ð²Ð½Ð¾ Ð¸Ð´ÐµÑ. Ð’ÑÑ‘ Ð½Ð°Ñ…Ð¾Ð´Ð¸Ñ‚ÑÑ Ð½Ðµ Ð² Ñ†Ð¸ÐºÐ»Ðµ, Ð´Ð»Ñ ÑƒÐ´Ð¾Ð±ÑÑ‚Ð²Ð° Ð´Ð°Ð»ÑŒÐ½ÐµÐ¹ÑˆÐµÐ¹ Ñ€Ð°Ð±Ð¾Ñ‚Ñ‹ Ñ Ð´Ð°Ð½Ð½Ð¾Ð¹ Ñ„ÑƒÐ½ÐºÑ†Ð¸ÐµÐ¹
def takeTable(nameDoc):
    # try:
    #     document = Document("C:/Users/UzziMauzer/Desktop/planprod/mvd/"+nameDoc)
    # except:
    #     print(nameDoc)
    #     subprocess.call(['soffice', '--headless', '--convert-to', 'docx', "C:/Users/UzziMauzer/Desktop/planprod/mvd/"+nameDoc])
    #     print("1")
    # document = Document("C:/Users/UzziMauzer/Desktop/planprod/"+nameDoc)

    #document = Document("/home/andrey/Desktop/planprod/"+nameDoc)
    # document= Document('C:/Users/UzziMauzer/Downloads/planqwer.docx')
    #document = Document('//home/user/planprod/mvd/plan/' + nameDoc)
    #document = Document('//root/planprod/mvd/' + nameDoc)
    document = Document("C:/Users/Andrey/Desktop/planprod/mvd/" + nameDoc)
    #document = Document('C:/Users/UzziMauzer/Desktop/planprod/mvd/' + nameDoc)

    # document = nameDoc
    listAllTable = []
    listOneTable = []
    listRow = []
    swichFlag = False
    table = document.tables
    print(len(table))
    if len(table) != 18:
        print('ne 18')
        if("Москва" in document.sections[0].first_page_footer.paragraphs[0].text):
            return (tekeUpdateTable(document))
        if len(table) != 19:
            print('ne 19')
            print(nameDoc)
            return ('dolbaeb')
        else:
            switchFlag = True

    #Ð”Ð»Ñ Ð¿ÐµÑ€Ð²Ð¾Ð¹ Ñ‚Ð°Ð±Ð»Ð¸Ñ†Ñ‹
    '''
    table = document.tables[5]
    countRow = 0
    for row in (table.rows):
        listRow = []
        if countRow > 0:
            countColumn = 0
            for cell in row.cells:
                if countRow != 15:
                    listRow.append(cell.text)
                else:
                    if countColumn > 0:
                        listRow.append(cell.text)
                countColumn+=1
            listOneTable.append(listRow)
        countRow+=1
    listAllTable.append(listOneTable)
    listOneTable = []

    #Ð”Ð»Ñ Ð²Ñ‚Ð¾Ñ€Ð¾Ð¹ Ñ‚Ð°Ð±Ð»Ð¸Ñ†Ñ‹
    table = document.tables[6]
    countRow = 0
    for row in (table.rows):
        listRow = []
        if countRow > 0:
            countColumn = 0
            for cell in row.cells:
                if countRow < 14:
                    listRow.append(cell.text)
                else:
                    if countColumn > 0:
                        listRow.append(cell.text)
                countColumn+=1
            listOneTable.append(listRow)
        countRow+=1
    listAllTable.append(listOneTable)
    listOneTable = []

    #Ð”Ð»Ñ 3Ð¹ Ñ‚Ð°Ð±Ð»Ð¸Ñ†Ñ‹
    table = document.tables[7]
    countRow = 0
    for row in (table.rows):
        listRow = []
        if countRow > 0:
            countColumn = 0
            for cell in row.cells:
                if countColumn > 0:
                    listRow.append(cell.text)
                countColumn+=1
            listOneTable.append(listRow)
        countRow+=1
    listAllTable.append(listOneTable)
    listOneTable = []

    #Ð”Ð»Ñ 4Ð¹ Ñ‚Ð°Ð±Ð»Ð¸Ñ†Ñ‹
    table = document.tables[8]
    countRow = 0
    for row in (table.rows):
        listRow = []
        if countRow > 0:
            for cell in row.cells:
                listRow.append(cell.text)
            listOneTable.append(listRow)
        countRow+=1
    listAllTable.append(listOneTable)
    listOneTable = []

    #Ð”Ð»Ñ 5Ð¹ Ñ‚Ð°Ð±Ð»Ð¸Ñ†Ñ‹
    table = document.tables[9]
    countRow = 0
    for row in (table.rows):
        listRow = []
        if countRow > 0:
            for cell in row.cells:
                 listRow.append(cell.text)
            listOneTable.append(listRow)
        countRow+=1
    listAllTable.append(listOneTable)
    listOneTable = []
    '''

    flag = False

    countLoose = 0

    #Ежемас Таблица
    table = document.tables[6]
    countRow = 0
    for row in (table.rows):
        listRow = []
        if countRow > 0:
            for cell in row.cells:
                if cell.text == '' or cell.text == ' ':
                    listRow.append('0')
                else:
                    cell.text = cell.text.replace(',','.')
                    listRow.append(cell.text)
            listOneTable.append(listRow)
        countRow+=1
    listAllTable.append(listOneTable)
    listOneTable = []

    #Для 6й таблицы
    table = document.tables[7]
    countRow = 0
    for row in (table.rows):
        listRow = []
        if countRow > 0:
            for cell in row.cells:
                if cell.text == '':
                    countLoose += 1
                if countLoose > 2:
                    flag = True
            countLoose = 0
            if flag == True:
                flag = False
                continue
            for cell in row.cells:
                listRow.append(cell.text)
            listOneTable.append(listRow)
        countRow+=1
    listAllTable.append(listOneTable)
    listOneTable = []

    #Для 7й таблицы
    table = document.tables[8]
    countRow = 0
    for row in (table.rows):
        listRow = []
        if countRow > 0:
            for cell in row.cells:
                if cell.text == '':
                    countLoose += 1
                if countLoose > 2:
                    flag = True
            countLoose = 0
            if flag == True:
                flag = False
                continue
            for cell in row.cells:
                listRow.append(cell.text)
            listOneTable.append(listRow)
        countRow+=1
    listAllTable.append(listOneTable)
    listOneTable = []

    #Для 8й таблицы
    table = document.tables[9]
    countRow = 0
    for row in (table.rows):
        listRow = []
        if countRow > 0:
            for cell in row.cells:
                if cell.text == '':
                    countLoose += 1
                if countLoose > 2:
                    flag = True
            countLoose = 0
            if flag == True:
                flag = False
                continue
            for cell in row.cells:
                listRow.append(cell.text)
            listOneTable.append(listRow)
        countRow+=1
    listAllTable.append(listOneTable)
    listOneTable = []

    #Для 9й таблицы
    table = document.tables[10]
    countRow = 0
    for row in (table.rows):
        listRow = []
        if countRow > 0:
            for cell in row.cells:
                if cell.text == '':
                    countLoose += 1
                if countLoose > 2:
                    flag = True
            countLoose = 0
            if flag == True:
                flag = False
                continue
            for cell in row.cells:
                listRow.append(cell.text)
            listOneTable.append(listRow)
        countRow+=1
    listAllTable.append(listOneTable)
    listOneTable = []













    #Для 5й 2й таблицы
    table = document.tables[11]
    countRow = 0
    for row in (table.rows):
        listRow = []
        if countRow > 1:
            for cell in row.cells:
                #Ищем пустые строки
                if cell.text == '':
                    countLoose += 1
                if countLoose > 2:
                    flag = True
                #Ищем 2е полугодие, если нет пустых строк
                if '2 полугодие' in cell.text.lower():
                    flag = True
                    break
            countLoose = 0
            if flag == True:
                flag = False
                break

            for cell in row.cells:
                listRow.append(cell.text)
            listOneTable.append(listRow)
        countRow+=1
    listAllTable.append(listOneTable)
    listOneTable = []

    countRow = 0
    for row in (table.rows):
        listRow = []
        if countRow > 1:
            for cell in row.cells:
                if '2 полугодие' in cell.text.lower():
                    flag = True
                    break
            if flag == True:
                flag = False
                break
        countRow+=1


    needCountRow = countRow



    countRow=0
    for row in (table.rows):
        listRow = []
        if countRow > needCountRow:


            for cell in row.cells:
                if cell.text == '':
                    countLoose += 1
                if countLoose > 2:
                    flag = True
            countLoose = 0
            if flag == True:
                flag = False
                break
            for cell in row.cells:
                listRow.append(cell.text)
            listOneTable.append(listRow)
        countRow+=1
    listAllTable.append(listOneTable)
    listOneTable = []



    #Для 6й 2й таблицы
    table = document.tables[12]
    countRow = 0
    for row in (table.rows):
        listRow = []
        if countRow > 1:
            for cell in row.cells:
                #Ищем пустые строки
                if cell.text == '':
                    countLoose += 1
                if countLoose > 2:
                    flag = True
                #Ищем 2е полугодие, если нет пустых строк
                if '2 полугодие' in cell.text.lower():
                    flag = True
                    break
            countLoose = 0
            if flag == True:
                flag = False
                break
            for cell in row.cells:
                listRow.append(cell.text)
            listOneTable.append(listRow)
        countRow+=1
    listAllTable.append(listOneTable)
    listOneTable = []


    countRow = 0
    for row in (table.rows):
        listRow = []
        if countRow > 1:
            for cell in row.cells:
                if '2 полугодие' in cell.text.lower():
                    flag = True
                    break
            if flag == True:
                flag = False
                break
        countRow+=1


    needCountRow = countRow




    countRow=0
    for row in (table.rows):
        listRow = []
        if countRow > needCountRow:


            for cell in row.cells:
                if cell.text == '':
                    countLoose += 1
                if countLoose > 2:
                    flag = True
            countLoose = 0
            if flag == True:
                flag = False
                break
            for cell in row.cells:
                listRow.append(cell.text)
            listOneTable.append(listRow)
        countRow+=1
    listAllTable.append(listOneTable)
    listOneTable = []



    if swichFlag == False:
        #Для 7й 2й таблицы
        table = document.tables[13]
        countRow = 0
        for row in (table.rows):
            listRow = []
            if countRow > 1:
                for cell in row.cells:
                    #Ищем пустые строки
                    if cell.text == '':
                        countLoose += 1
                    if countLoose > 2:
                        flag = True
                    #Ищем 2е полугодие, если нет пустых строк
                    if '2 полугодие' in cell.text.lower():
                        flag = True
                        break
                countLoose = 0
                if flag == True:
                    flag = False
                    break
                for cell in row.cells:
                    listRow.append(cell.text)
                listOneTable.append(listRow)
            countRow+=1
        listAllTable.append(listOneTable)
        listOneTable = []


        countRow = 0
        for row in (table.rows):
            listRow = []
            if countRow > 1:
                for cell in row.cells:
                    if '2 полугодие' in cell.text.lower():
                        flag = True
                        break
                if flag == True:
                    flag = False
                    break
            countRow+=1


        needCountRow = countRow




        countRow=0
        for row in (table.rows):
            listRow = []
            if countRow > needCountRow:


                for cell in row.cells:
                    if cell.text == '':
                        countLoose += 1
                    if countLoose > 2:
                        flag = True
                countLoose = 0
                if flag == True:
                    flag = False
                    break
                for cell in row.cells:
                    listRow.append(cell.text)
                listOneTable.append(listRow)
            countRow+=1
        listAllTable.append(listOneTable)
        listOneTable = []
    else:
        table = document.tables[13]
        countRow = 0
        for row in (table.rows):
            listRow = []
            if countRow > 0:
                for cell in row.cells:
                    if cell.text == '':
                        countLoose += 1
                    if countLoose > 2:
                        flag = True
                countLoose = 0
                if flag == True:
                    flag = False
                    continue
                for cell in row.cells:
                    listRow.append(cell.text)
                listOneTable.append(listRow)
            countRow+=1
        listAllTable.append(listOneTable)
        listOneTable = []



        table = document.tables[14]
        countRow = 0
        for row in (table.rows):
            listRow = []
            if countRow > 0:
                for cell in row.cells:
                    if cell.text == '':
                        countLoose += 1
                    if countLoose > 2:
                        flag = True
                countLoose = 0
                if flag == True:
                    flag = False
                    continue
                for cell in row.cells:
                    listRow.append(cell.text)
                listOneTable.append(listRow)
            countRow+=1
        listAllTable.append(listOneTable)
        listOneTable = []













    '''
    #Ð”Ð»Ñ 11Ð¹ 2Ð¹ Ñ‚Ð°Ð±Ð»Ð¸Ñ†Ñ‹
    table = document.tables[11]
    countRow = 0
    for row in (table.rows):
        listRow = []
        if countRow > 1:
            for cell in row.cells:
                if cell.text != '':
                    break
                else:
                    flag = True
            if flag == True:
                flag = False
                break
            for cell in row.cells:
                listRow.append(cell.text)
            listOneTable.append(listRow)
        countRow+=1
    listAllTable.append(listOneTable)
    listOneTable = []

    table = document.tables[11]
    for row in (table.rows):
        listRow = []
        if countRow > 1:
            for cell in row.cells:
                if '2 Ð¿Ð¾Ð»ÑƒÐ³Ð¾Ð´Ð¸Ðµ' in cell.text.lower():
                    flag = True
                    break
            if flag == True:
                flag = False
                break
        countRow+=1


    needCountRow = countRow - 3


    table = document.tables[11]
    countRow=0
    for row in (table.rows):
        listRow = []
        if countRow > needCountRow:

            for cell in row.cells:


            for cell in row.cells:
                if cell.text != '':
                    break
                else:
                    flag = True
            if flag == True:
                flag = False
                break
            for cell in row.cells:
                listRow.append(cell.text)
            listOneTable.append(listRow)
        countRow+=1
    listAllTable.append(listOneTable)
    listOneTable = []



    #Ð”Ð»Ñ 12Ð¹ 2Ð¹ Ñ‚Ð°Ð±Ð»Ð¸Ñ†Ñ‹
    table = document.tables[12]
    countRow = 0
    for row in (table.rows):
        listRow = []
        if countRow > 1:
            for cell in row.cells:
                if cell.text != '':
                    break
                else:
                    flag = True
            if flag == True:
                flag = False
                break
            for cell in row.cells:
                listRow.append(cell.text)
            listOneTable.append(listRow)
        countRow+=1
    listAllTable.append(listOneTable)
    listOneTable = []

    table = document.tables[12]
    for row in (table.rows):
        listRow = []
        if countRow > 1:
            for cell in row.cells:
                if '2 Ð¿Ð¾Ð»ÑƒÐ³Ð¾Ð´Ð¸Ðµ' in cell.text.lower():
                    flag = True
                    break
            if flag == True:
                flag = False
                break
        countRow+=1


    needCountRow = countRow -2


    table = document.tables[12]
    countRow=0
    for row in (table.rows):
        listRow = []
        if countRow > needCountRow:

            for cell in row.cells:


            for cell in row.cells:
                if cell.text != '':
                    break
                else:
                    flag = True
            if flag == True:
                flag = False
                break
            for cell in row.cells:
                listRow.append(cell.text)
            listOneTable.append(listRow)
        countRow+=1
    listAllTable.append(listOneTable)
    listOneTable = []



    #Ð”Ð»Ñ 13Ð¹ 2Ð¹ Ñ‚Ð°Ð±Ð»Ð¸Ñ†Ñ‹
    table = document.tables[13]
    countRow = 0
    for row in (table.rows):
        listRow = []
        if countRow > 1:
            for cell in row.cells:
                if cell.text != '':
                    break
                else:
                    flag = True
            if flag == True:
                flag = False
                break
            for cell in row.cells:
                listRow.append(cell.text)
            listOneTable.append(listRow)
        countRow+=1
    listAllTable.append(listOneTable)
    listOneTable = []

    table = document.tables[13]
    for row in (table.rows):
        listRow = []
        if countRow > 1:
            for cell in row.cells:
                if '2 Ð¿Ð¾Ð»ÑƒÐ³Ð¾Ð´Ð¸Ðµ' in cell.text.lower():
                    flag = True
                    break
            if flag == True:
                flag = False
                break
        countRow+=1


    needCountRow = countRow -5


    table = document.tables[13]
    countRow=0
    for row in (table.rows):
        listRow = []
        if countRow > needCountRow:

            for cell in row.cells:


            for cell in row.cells:
                if cell.text != '':
                    break
                else:
                    flag = True
            if flag == True:
                flag = False
                break
            for cell in row.cells:
                listRow.append(cell.text)
            listOneTable.append(listRow)
        countRow+=1
    listAllTable.append(listOneTable)
    listOneTable = []
    '''


    '''


    #Ð”Ð»Ñ 10Ð¹ Ñ‚Ð°Ð±Ð»Ð¸Ñ†Ñ‹
    table = document.tables[11]
    countRow = 0
    for row in (table.rows):
        listRow = []
        if countRow > 1 and countRow < 23:
            countColumn = 0
            for cell in row.cells:
                if countRow < 23:
                    listRow.append(cell.text)
            listOneTable.append(listRow)
        countRow+=1
    listAllTable.append(listOneTable)
    listOneTable = []

    #Ð”Ð»Ñ 11Ð¹ Ñ‚Ð°Ð±Ð»Ð¸Ñ†Ñ‹
    table = document.tables[11]
    countRow = 0
    for row in (table.rows):
        listRow = []
        if countRow > 24:
            countColumn = 0
            for cell in row.cells:
                if countRow > 24:
                    listRow.append(cell.text)
            listOneTable.append(listRow)
        countRow+=1
    listAllTable.append(listOneTable)
    listOneTable = []

    #Ð”Ð»Ñ 12Ð¹ Ñ‚Ð°Ð±Ð»Ð¸Ñ†Ñ‹
    table = document.tables[12]
    countRow = 0
    for row in (table.rows):
        listRow = []
        if countRow > 1 and countRow < 24:
            countColumn = 0
            for cell in row.cells:
                if countRow < 24:
                    listRow.append(cell.text)
            listOneTable.append(listRow)
        countRow+=1
    listAllTable.append(listOneTable)
    listOneTable = []

    #Ð”Ð»Ñ 13Ð¹ Ñ‚Ð°Ð±Ð»Ð¸Ñ†Ñ‹
    table = document.tables[12]
    countRow = 0
    for row in (table.rows):
        listRow = []
        if countRow > 0 and countRow > 25:
            countColumn = 0
            for cell in row.cells:
                if countRow > 25:
                    listRow.append(cell.text)
            listOneTable.append(listRow)
        countRow+=1
    listAllTable.append(listOneTable)
    listOneTable = []

    #Ð”Ð»Ñ 14Ð¹ Ñ‚Ð°Ð±Ð»Ð¸Ñ†Ñ‹
    table = document.tables[13]
    countRow = 0
    for row in (table.rows):
        listRow = []
        if countRow > 0:
            for cell in row.cells:
                listRow.append(cell.text)
            listOneTable.append(listRow)
        countRow+=1
    listAllTable.append(listOneTable)
    listOneTable = []
    '''
    return listAllTable


def checkPrepods(nameDoc, namePrepod):
    wb = openpyxl.load_workbook(filename=nameDoc, data_only=True)

    sheets = wb.sheetnames
    ws = wb[sheets[0]]
    prepodRow = 1

    # Находим строку с преподом
    for column in ws.iter_rows(min_col=2, min_row=1, max_col=2, max_row=1000, values_only=True):
        if str(namePrepod).lower() in str(column[0]).lower():
            break
        prepodRow += 1

    if prepodRow == 1001:
        return (
                   "Произошла ошибка: преподаватель %s не был найден в EXCEL таблице, пожалуйста, проверьте, что столбец с преподавателями находится в столбце 'B'. Иначе проверьте правописание фамилии в документе на орфографические ошибки.") % namePrepod

    secondPrepod = 1
    for column in ws.iter_rows(min_col=2, min_row=(prepodRow + 1), max_col=2, max_row=(prepodRow + 30),
                               values_only=True):

        reg = re.search("^([А-ЯЁ].+?\s)[А-ЯЁ]\.[А-ЯЁ]\.", str(column[0]))
        if reg and len(reg.group(1)) > 0:
            count = secondPrepod
            break
        reg = re.search("(ИТОГО):", str(column[0]))
        if reg and len(reg.group(1)) > 0:
            count = secondPrepod
            break

        secondPrepod += 1

    if secondPrepod == 1:
        return (
            'Произошла ошибка: в вашем документе 2 фамилии стоят в двух строках подряд, нельзя писать фамилии преподавателей в строках с названиями дискиплин, если они идут по замене.')

    return 'Преподаватель %s найден' % namePrepod

def checkDocumentXLS(nameDoc, flag):
    wb = openpyxl.load_workbook(filename =nameDoc, data_only=True)
    if flag==True:
        zapasFact=[]
        zapasFactAudit=[]
        nameNeedColunm = ['лекции', 'семинары', 'практические занятия в группе', 'практические занятия в подгруппе', 'круглый стол', 'консультации перед экзаменами', 'текущие консультации', 'внеаудиторное чтение', 'практика руководство', 'ВКР', 'курсовая работа', 'контрольная работа аудиторная', 'контрольная работа домашняя', 'проверка практикума', 'проверка лабораторной работы', 'защита практики', 'зачет устный', 'зачет письменный', 'вступительные испытания', 'экзамены', 'государственные экзамены', 'юнктура)', 'руководство адъюнктами', 'Плановая нагрузка', 'Плановая аудиторная нагрузка']
    else:
        nameNeedColunm = ['лекции', 'семинары', 'практические занятия в группе', 'практические занятия в подгруппе', 'круглый стол', 'консультации перед экзаменами', 'текущие консультации', 'внеаудиторное чтение', 'практика руководство', 'ВКР', 'курсовая работа', 'контрольная работа аудиторная', 'контрольная работа домашняя', 'проверка практикума', 'проверка лабораторной работы', 'защита практики', 'зачет устный', 'зачет письменный', 'вступительные испытания', 'экзамены', 'государственные экзамены', 'юнктура)', 'руководство адъюнктами', 'Фактически выполненная нагрузка', 'Фактически выполненная аудиторная нагрузка']
        zapasFact=['Фактически выполненная нагрузка', 'выполненная нагрузка','выполненая нагрузка','Фактическая нагрузка']
        zapasFactAudit=['Фактически выполненная аудиторная нагрузка', 'выполненная аудиторная нагрузка','выполненая аудиторная нагрузка','Фактическая аудиторная нагрузка']
    nameprepodColumn = ['наименование дисциплины', 'Факультет', 'Курс, группа']
    listFirst = []
    listSecond = []
    listResult = []
    needColumnFirst = []
    needColumnSecond = []
    needColumnThird = []
    needPrepodColumn = []
    sheets = wb.sheetnames
    ws = wb[sheets[0]]
    prepodRow = 1
    #Защита от дураков на ошибки
    for textCol in zapasFact:
        for column in ws.iter_cols(min_col=1, min_row=8, max_col=500, max_row=8):
            if textCol.lower() in str(column[0].value).lower():
                column[0].value = 'фактически выполненная нагрузка'
    for textCol in zapasFactAudit:
        for column in ws.iter_cols(min_col=1, min_row=8, max_col=500, max_row=8):
                if textCol.lower() in str(column[0].value).lower():
                    column[0].value = 'фактически выполненная аудиторная нагрузка'
    #Функция заполнения Индексов столбцов
    for textCol in nameNeedColunm:
        k = 0
        for column in ws.iter_cols(min_col=1, min_row=8, max_col=200, max_row=8):
                if textCol.lower() in str(column[0].value).lower():
                    if textCol.lower() == 'экзамены':
                        if textCol.lower() != str(column[0].value).lower():
                            continue
                    k += 1
                    column[0].value = ' '
                    reg = re.search(".+?\.(.+?)8>", str(column[0]))
                    if reg and len(reg.group(1)) > 0:
                        needText = reg.group(1)
                    if k == 1:
                        needColumnFirst.append(needText)
                    elif k == 2:
                        needColumnSecond.append(needText)
                    else: needColumnThird.append(needText)
        if k != 3 :
            if textCol == 'лекции':
                return ('Произошла ошибка: пожалуйста, проверьте, что в вашем EXCEL документе только один лист (одна вкладка в нижнем левом углу), если это так, то либо шапка таблицы с названиями видов занятий находится не на 8й строке, либо орфографическая ошибка в слове "лекции"')
            elif textCol == 'фактически выполненная нагрузка':
                return ('Произошла ошибка: пожалуйста, проверьте, что в вашем EXCEL документе только один лист (одна вкладка в нижнем левом углу), если это так, то возможна ошибка в слове "фактически выполненная нагрузка" (Чтобы убедиться в этом, вам нужно нажать CTRL+F в документе, в поле найти ввести "фактически выполненная нагрузка" и нажать "найти все", в выводе должно отобразиться 3 строки)')
            return ('Произошла ошибка: пожалуйста, проверьте, что в вашем EXCEL документе только один лист (одна вкладка в нижнем левом углу), если это так, то возможна ошибка в слове %s (Чтобы убедиться в этом, вам нужно нажать CTRL+F в документе, в поле найти ввести "%s" и нажать "найти все", в выводе должно отобразиться 3 строки)') % (textCol, textCol)
    for textCol in nameprepodColumn:
        k = 0
        for column in ws.iter_cols(min_col=1, min_row=8, max_col=200, max_row=8):
                if textCol.lower() == str(column[0].value).lower():
                    k += 1
                    column[0].value = ' '
                    reg = re.search(".+?\.(.+?)8>", str(column[0]))
                    if reg and len(reg.group(1)) > 0:
                        needText = reg.group(1)
                    needPrepodColumn.append(needText)
        if k != 3:
            print(k)
            return ('Ошибка в одном из слов:' + str(', '.join(nameprepodColumn)))
    return 'Ваш документ соотвествует образцу'



if __name__ == '__main__':
    print("start")
    # xlsPrepod('322.xlsx')
#     fopen()
#     createDoc('demo', listEny)
    # takeTable('demo.docx')
#     readInfoDoc("demo.docx")
#     writeInfoDoc("test", listInfo)
#     takeXls('templateXls','Иванов В.Ю.')

'''
Установка дополнительных библиотек: pip install python-docx
Версия библиотеки: 0.8.10
Зависимости: Python 2.6, 2.7, 3.3, or 3.4     lxml >= 2.3.2
Документация: https://python-docx.readthedocs.io/en/latest/
'''
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from django.utils.lorem_ipsum import paragraph, paragraphs
from docx.shared import Pt

import openpyxl

# Функция для получения упорядоченного списка всех нужных переменных из xls файла и его возврат
def takeXls(nameDoc,namePrepod):
    wb = openpyxl.load_workbook(filename =nameDoc, data_only=True)

    needColumnFirst = ['O','Q','S','U','W','X','Y','AA','AC','AE','AG','AI','AK','AM','AO','AQ','AS','AU','AW','AY','BA','BC','BE','BF','BG' ]
    needColumnSecond = ['BZ','CA','CB','CD','CF','CH','CI','CJ','CL','CN','CP','CR','CT','CV','CX','CZ','DB','DD','DF','DH','DJ','DL','DN','DP','DQ','DR']
    needColumnThird = ['EK','EM','EO','EQ','ES','ET','EU','EW','EY','FA','FC','FE','FG','FI','FK','FM','FO','FQ','FS','FU','FW','FY','GA','GB','GC']
    listFirst = []
    listSecond = []
    listResult = []

    sheets = wb.sheetnames
    ws = wb[sheets[0]]
    prepodRow = 1
    #Находим строку с преподом

    for column in ws.iter_rows(min_col=2, min_row=1, max_col=2, max_row=400, values_only=True):
        if str(namePrepod).lower() in str(column[0]).lower():
            break
        prepodRow += 1

    #Получаем значения 1 таблицы в лист


    for takeStr in range(1,14):
        listFirst.append([])
        firstBigSell = str(ws['B'+str(prepodRow+takeStr)].value) +' '+ str(ws['E'+str(prepodRow+takeStr)].value) +' '+ str(ws['F'+str(prepodRow+takeStr)].value)
        if firstBigSell != 'None None None':
            listFirst[takeStr-1].append(firstBigSell)
        else:
            listFirst[takeStr-1].append('0')
        for takeColumn in needColumnFirst:
            cell = takeColumn+str(prepodRow+takeStr)
            if ws[cell].value == None or ws[cell].value == '0':
                listFirst[takeStr-1].append('0')
            else:
                listFirst[takeStr-1].append(ws[cell].value)
    listFirst.append([])
    for takeColumn in needColumnFirst:
            cell = takeColumn+str(prepodRow)
            if ws[cell].value == None or ws[cell].value == '0':
                listFirst[takeStr].append('0')
            else:
                listFirst[takeStr].append(ws[cell].value)



    #Получаем значения 2 таблицы в лист
    for takeStr in range(1,13):
        listSecond.append([])
        firstBigSell = str(ws['BM'+str(prepodRow+takeStr)].value) +' '+ str(ws['BP'+str(prepodRow+takeStr)].value) +' '+ str(ws['BQ'+str(prepodRow+takeStr)].value)
        if firstBigSell != 'None None None':
            listSecond[takeStr-1].append(firstBigSell)
        else:
            listSecond[takeStr-1].append('0')
        for takeColumn in needColumnSecond:
            cell = takeColumn+str(prepodRow+takeStr-1)
            if ws[cell].value == None or ws[cell].value == '0':
                listSecond[takeStr-1].append('0')
            else:
                listSecond[takeStr-1].append(ws[cell].value)
    listSecond.append([])
    for takeColumn in needColumnSecond:
            cell = takeColumn+str(prepodRow)
            if ws[cell].value == None or ws[cell].value == '0':
                listSecond[takeStr].append('0')
            else:
                listSecond[takeStr].append(ws[cell].value)
    listSecond.append([])
    #Получаем значения 3 таблицы в лист
    for takeColumn in needColumnThird:
            cell = takeColumn+str(prepodRow)
            if ws[cell].value == None or ws[cell].value == '0':
                listSecond[takeStr+1].append('0')
            else:
                listSecond[takeStr+1].append(ws[cell].value)

    listResult.append(listFirst)
    listResult.append(listSecond)
    return(listResult)

#/home/andrey/Documents/Vazhno_sho_pizdec/mvdproject/mvd/plan/
# Запись в документ Основной информации
def writeInfoDoc(nameDoc, listInfo):
    document = Document("test.docx")
    needStrings = [7,8,10,21,23,31]
    # Для строчек, который разделяются на 3, то есть кафедры
    threeStrings = [25,27,29]
    threeList = listInfo.pop(5)
    oneOfThreeList = [[],[],[]]
    numberSimbolAllStrings = 0


    countParagraphs = 0
    countListInfo = 0
    for paragraphs in document.paragraphs:
        if countParagraphs in needStrings:
            print(countListInfo)
            paragraphs.text = (listInfo[countListInfo])
            if(countParagraphs == 31 or countParagraphs == 23 or countParagraphs == 21):
                for run in paragraphs.runs:
                    font = run.font
                    font.bold = 1
            for run in paragraphs.runs:
                font = run.font
                font.size = Pt(14)
            countListInfo+=1

        # Для тех самых 3х строчек
        elif countParagraphs in threeStrings:
            for countStrings in range(0,3):
                if threeList != '':
                    if (len(threeList) - numberSimbolAllStrings) > 65:
                        flag = True
                        for indexSpase in range(65,0,-1):
                            if flag == True:
                                if threeList[indexSpase] == ' ':
                                    print('WOW')
                                    firstSpase = indexSpase
                                    numSimbol = 0
                                    while (numSimbol != firstSpase):
                                        oneOfThreeList[countStrings] += threeList[numSimbol + numberSimbolAllStrings]
                                        numSimbol +=1
                                    numberSimbolAllStrings += firstSpase + 1
                                    print(oneOfThreeList[countStrings])
                                    document.paragraphs[25+countStrings*2].text = (oneOfThreeList[countStrings])
                                    for run in document.paragraphs[25+countStrings*2].runs:
                                        font = run.font
                                        font.size = Pt(14)
                                    flag = False
                    else:
                        oneOfThreeList[countStrings] = threeList[numberSimbolAllStrings:]
                        document.paragraphs[threeStrings[countStrings]].text = (oneOfThreeList[countStrings])
                        for run in document.paragraphs[25+countStrings*2].runs:
                            font = run.font
                            font.size = Pt(14)
                        threeList =''
        countParagraphs+=1


    paragraph_format = document.paragraphs[10].paragraph_format
    paragraph_format.right_indent = Pt(55)
    document.save('/home/andrey/Documents/Vazhno_sho_pizdec/mvdproject/mvd/plan/' + nameDoc)


# Чтение из документа основной информации
def readInfoDoc(nameDoc):
    document = Document(nameDoc)
    needStrings = [7,8,10,21,23,25,27,29,31]
    infoDoc = []

    for num in needStrings:
        if num == 27 or num == 28:
            infoDoc[5] += document.paragraphs[num].text
        else:
            infoDoc.append(document.paragraphs[num].text)
    print(infoDoc)


# Заполнение документа по полученным данным и сохранение документа
def createDoc(nameDoc, listCell):
    document = Document('/home/andrey/Documents/Vazhno_sho_pizdec/mvdproject/mvd/plan/test.docx')

    count = 0
    for table in document.tables:

        for row in (table.rows):
            for cell in row.cells:
                if cell.text.lower() == 'x':
                    cell.text = str(listCell[count])
                    cell.alignment=WD_ALIGN_PARAGRAPH.CENTER
                    # cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    count+=1


    # document.save('/home/andrey/Documents/Vazhno_sho_pizdec/mvdproject/mvd/plan/' + nameDoc + '.docx')
    return document

# В данной функции мы Обрабатываем получаемый документ и возвращаем все таблицы во вложенных списках
# Советую тут ничего не менять, говно идея. Всё находится не в цикле, для удобства дальнейшей работы с данной функцией
def takeTable(nameDoc):
    document = Document(nameDoc)
    listAllTable = []
    listOneTable = []
    listRow = []

    # Для превой таблицы
    table = document.tables[2]
    countRow = 0
    for row in (table.rows):
        listRow = []
        if countRow > 2:
            countColumn = 0
            for cell in row.cells:
                if countColumn > 0:
                    listRow.append(cell.text)
                countColumn+=1
            listOneTable.append(listRow)
        countRow+=1
    listAllTable.append(listOneTable)
    listOneTable = []

    # Для второй таблицы
    table = document.tables[5]
    countRow = 0
    for row in (table.rows):
        listRow = []
        if countRow > 0:
            countColumn = 0
            for cell in row.cells:
                if countRow != 12:
                    listRow.append(cell.text)
                else:
                    if countColumn > 0:
                        listRow.append(cell.text)
                countColumn+=1
            listOneTable.append(listRow)
        countRow+=1
    listAllTable.append(listOneTable)
    listOneTable = []

    # Для 3 таблицы
    table = document.tables[6]
    countRow = 0
    for row in (table.rows):
        listRow = []
        if countRow > -1:
            countColumn = 0
            for cell in row.cells:
                if countRow < 11:
                    listRow.append(cell.text)
                else:
                    if countColumn > 0:
                        listRow.append(cell.text)
                countColumn+=1
            listOneTable.append(listRow)
        countRow+=1
    listAllTable.append(listOneTable)
    listOneTable = []

    # Для 4 таблицы
    table = document.tables[7]
    countRow = 0
    for row in (table.rows):
        listRow = []
        if countRow > 0:
            for cell in row.cells:
                listRow.append(cell.text)
            listOneTable.append(listRow)
        countRow+=1
    listAllTable.append(listOneTable)
    listOneTable = []

    # Для 5 таблицы
    table = document.tables[8]
    countRow = 0
    for row in (table.rows):
        listRow = []
        if countRow > -1:
            for cell in row.cells:
                listRow.append(cell.text)
            listOneTable.append(listRow)
        countRow+=1
    listAllTable.append(listOneTable)
    listOneTable = []

    # Для 6 таблицы
    table = document.tables[9]
    countRow = 0
    for row in (table.rows):
        listRow = []
        if countRow > 0:
            countColumn = 0
            for cell in row.cells:
                if countRow != 14:
                    listRow.append(cell.text)
                else:
                    if countColumn > 0:
                        listRow.append(cell.text)
                countColumn+=1
            listOneTable.append(listRow)
        countRow+=1
    listAllTable.append(listOneTable)
    listOneTable = []

    # Для 7 таблицы
    table = document.tables[10]
    countRow = 0
    for row in (table.rows):
        listRow = []
        if countRow > -1:
            countColumn = 0
            for cell in row.cells:
                if countRow != 5:
                    listRow.append(cell.text)
                else:
                    if countColumn > 0:
                        listRow.append(cell.text)
                countColumn+=1
            listOneTable.append(listRow)
        countRow+=1
    listAllTable.append(listOneTable)
    listOneTable = []

    # Для 8 таблицы
    table = document.tables[11]
    countRow = 0
    for row in (table.rows):
        listRow = []
        if countRow > 0:
            for cell in row.cells:
                listRow.append(cell.text)
            listOneTable.append(listRow)
        countRow+=1
    listAllTable.append(listOneTable)
    listOneTable = []

    # Для 9 таблицы
    table = document.tables[12]
    countRow = 0
    for row in (table.rows):
        listRow = []
        if countRow > -1:
            for cell in row.cells:
                listRow.append(cell.text)
            listOneTable.append(listRow)
        countRow+=1
    listAllTable.append(listOneTable)
    listOneTable = []

    # Для 10 таблицы
    table = document.tables[13]
    countRow = 0
    for row in (table.rows):
        listRow = []
        if countRow > 0:
            countColumn = 0
            for cell in row.cells:
                if countRow != 14:
                    listRow.append(cell.text)
                else:
                    if countColumn > 0:
                        listRow.append(cell.text)
                countColumn+=1
            listOneTable.append(listRow)
        countRow+=1
    listAllTable.append(listOneTable)
    listOneTable = []

    # Для 11 таблицы
    table = document.tables[14]
    countRow = 0
    for row in (table.rows):
        listRow = []
        if countRow > -1:
            countColumn = 0
            for cell in row.cells:
                if countRow < 8:
                    if countColumn > 0:
                        listRow.append(cell.text)
                else:
                    if countColumn > 0:
                        listRow.append(cell.text)
                countColumn+=1
            listOneTable.append(listRow)
        countRow+=1
    listAllTable.append(listOneTable)
    listOneTable = []

    # Для 12 таблицы
    table = document.tables[15]
    countRow = 0
    for row in (table.rows):
        listRow = []
        if countRow > 0:
            for cell in row.cells:
                listRow.append(cell.text)
            listOneTable.append(listRow)
        countRow+=1
    listAllTable.append(listOneTable)
    listOneTable = []


    # Для 13 таблицы
    table = document.tables[16]
    countRow = 0
    for row in (table.rows):
        listRow = []
        if countRow > -1:
            for cell in row.cells:
                listRow.append(cell.text)
            listOneTable.append(listRow)
        countRow+=1
    listAllTable.append(listOneTable)
    listOneTable = []

    # Для 14 таблицы
    table = document.tables[17]
    countRow = 0
    for row in (table.rows):
        listRow = []
        if countRow > 1:
            countColumn = 0
            for cell in row.cells:
                if countRow != 15:
                    listRow.append(cell.text)
                else:
                    if countColumn > 2:
                        listRow.append(cell.text)
                countColumn+=1
            listOneTable.append(listRow)
        countRow+=1
    listAllTable.append(listOneTable)
    listOneTable = []

    # Для 15 таблицы
    table = document.tables[18]
    countRow = 0
    for row in (table.rows):
        listRow = []
        if countRow > 1:
            countColumn = 0
            for cell in row.cells:
                if countRow < 22:
                    listRow.append(cell.text)
                else:
                    if countColumn > 2:
                        listRow.append(cell.text)
                countColumn+=1
            listOneTable.append(listRow)
        countRow+=1
    listAllTable.append(listOneTable)
    listOneTable = []

    # Для 16 таблицы
    table = document.tables[19]
    countRow = 0
    for row in (table.rows):
        listRow = []
        if countRow > 1:
            countColumn = 0
            for cell in row.cells:
                if countRow != 21:
                    listRow.append(cell.text)
                else:
                    if countColumn > 2:
                        listRow.append(cell.text)
                countColumn+=1
            listOneTable.append(listRow)
        countRow+=1
    listAllTable.append(listOneTable)
    listOneTable = []

    # Для 17 таблицы
    table = document.tables[20]
    countRow = 0
    for row in (table.rows):
        listRow = []
        if countRow > 1:
            countColumn = 0
            for cell in row.cells:
                if countRow < 23:
                    listRow.append(cell.text)
                else:
                    if countColumn > 2:
                        listRow.append(cell.text)
                countColumn+=1
            listOneTable.append(listRow)
        countRow+=1
    listAllTable.append(listOneTable)
    listOneTable = []

    # Для 18 таблицы
    table = document.tables[21]
    countRow = 0
    for row in (table.rows):
        listRow = []
        if countRow > 1:
            countColumn = 0
            for cell in row.cells:
                if countRow != 15 and countRow != 16 and countRow != 30 and countRow != 31:
                    listRow.append(cell.text)
                else:
                    if countRow == 15 or countRow == 30 or countRow == 31:
                        if countColumn > 1:
                            listRow.append(cell.text)
                countColumn+=1
            if countRow != 16:
                listOneTable.append(listRow)
        countRow+=1
    listAllTable.append(listOneTable)
    listOneTable = []

    # Для 19 таблицы
    table = document.tables[22]
    countRow = 0
    for row in (table.rows):
        listRow = []
        if countRow > 1:
            countColumn = 0
            for cell in row.cells:
                if countRow != 10 and countRow != 11 and countRow != 21 and countRow != 22:
                    listRow.append(cell.text)
                else:
                    if countRow == 10 or countRow == 21 or countRow == 22:
                        if countColumn > 1:
                            listRow.append(cell.text)
                countColumn+=1
            if countRow != 11:
                listOneTable.append(listRow)
        countRow+=1
    listAllTable.append(listOneTable)
    # print(listAllTable)

    return listAllTable

if __name__ == '__main__':
    main()

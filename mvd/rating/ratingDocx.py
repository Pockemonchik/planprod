# -*- coding: utf8 -*-
'''
Установка дополнительных библиотек: pip install python-docx, pip install openxls
Версия библиотеки: 0.8.10
Зависимости: Python 2.6, 2.7, 3.3, or 3.4     lxml >= 2.3.2
Документация: https://python-docx.readthedocs.io/en/latest/
'''
from docx import Document
from docx.shared import Pt

def createRatingDocx(toptext, tableLens, inTable, sumBal):
    document = Document('/home/user/planprod/mvd/rating/RateShablon.docx')

    document.paragraphs[0].text="Рейтинговый информационный лист педагогической деятельности кафедры %s по итогам работы в %s/%s учебном году" % (toptext[0], toptext[1], toptext[2])
    document.paragraphs[1].text="Фамилия, имя, отчество: %s" % (toptext[3])
    document.paragraphs[2].text="Должность, ставка: %s %s" % (toptext[4], toptext[5])
    document.paragraphs[3].text="Учёная стеепнь, учёное звание: %s %s" % (toptext[6], toptext[7])
    document.paragraphs[4].text="Специальное звание: %s" % (toptext[8])

    for run in document.paragraphs[0].runs:
        font = run.font
        font.name= "Times New Roman"
        font.size = Pt(14)

    for i in range (4):
        for run in document.paragraphs[i+1].runs:
            font = run.font
            font.name= "Times New Roman"
            font.size = Pt(12)

    for i in range (4):
        table = document.tables[i]
        for j in range (tableLens[i]):
            table.add_row()
        count = 0
        for row in (table.rows):
            for cell in row.cells:
                if cell.text.lower() == '':
                    cell.text = inTable[i][count]
                    count += 1

    for i in range (4):
        table = document.tables[i]
        table.add_row()
        table.rows[-1].cells[-2].text = "Cуммарный балл по разделу %s" % (i+1)
        table.rows[-1].cells[-1].text = sumBal[i]

    table = document.tables[3]
    table.add_row()
    table.rows[-1].cells[-2].text = "Итоговый рейтинг"
    table.rows[-1].cells[-1].text = sumBal[-1]


    document.paragraphs[-2].text = "Начальник кафедры %s" % (toptext[0])
    for run in document.paragraphs[-2].runs:
        font = run.font
        font.name= "Times New Roman"
        font.size = Pt(12)

    document.paragraphs[-1].text = toptext[8] + " "*70 +  '{} {:.1}. {:.1}.'.format(*toptext[3].split())
    for run in document.paragraphs[-1].runs:
        font = run.font
        font.name= "Times New Roman"
        font.size = Pt(12)



    # document.save('C:\\Users\\UzziMauzer\\Desktop\\pythonDocx\\tmp.docx')
    return document

if __name__ == "__main__":
    toptext = ["ИБ", "2020", "2021", "Фамилия Имя Отчество", "Должность", "ставка", "Учёная стеепнь", "учёное звание", "Специальное звание"]
    tableLens = [5,3,1,2]
    inTable = [["ИБ", "2020", "2021", "Крот Пот Мот", "Пост мост", "Курица 1","ИБ", "2020", "2021", "Крот Пот Мот", "Пост мост", "Курица 1","ИБ", "2020", "2021"],
                ["ИБ", "2020", "2021", "Крот Пот Мот", "Пост мост", "Курица 1","ИБ", "2020", "2021"],
                ["ИБ", "2020", "2021"],
                ["ИБ", "2020", "2021", "Крот Пот Мот", "Пост мост", "Курица 1"]]
    sumBal = ["10","20","30","20","80"]
    createRatingDocx(toptext, tableLens, inTable, sumBal)
